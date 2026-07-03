import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================
# CONFIG — под проект CityPark
# =========================

PROJECT_NAME = "CityPark"
DEFAULT_PROJECT_DIR = "City__Park-main"
OUTPUT_FILE = "Listing_CityPark.docx"

ALLOWED_EXTENSIONS = {
    ".py", ".html", ".css", ".js", ".json", ".md"
}

IGNORED_DIRS = {
    ".git", "__pycache__", ".idea", ".vscode",
    ".venv", "venv", "env",
    "node_modules", "dist", "build",
    "media", "staticfiles",
}

# Миграции Django обычно генерируются автоматически и сильно раздувают листинг.
# Если преподаватель требует абсолютно все .py-файлы, поменяй на True.
INCLUDE_MIGRATIONS = False

MAX_FILE_SIZE_KB = 800
ADD_LINE_NUMBERS = False

# Порядок вывода файлов в листинге
APP_ORDER = [
    "manage.py",
    "CityPark",
    "users",
    "catalog",
    "restaurant",
    "templates",
    "static",
    "README.md",
]

# =========================
# WORD HELPERS
# =========================

def set_columns(section, count=2, space_twips="360"):
    """Настройка количества колонок в секции Word."""
    sect_pr = section._sectPr

    for child in list(sect_pr.findall(qn("w:cols"))):
        sect_pr.remove(child)

    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))
    sect_pr.append(cols)


def set_page_layout(section):
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)


def set_default_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(10)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = 1
    run = paragraph.add_run()

    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def clean_text(text):
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)


def add_small_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(8.5)
    return p


def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    run.font.size = Pt(7.5)


def add_code_block(doc, content):
    content = clean_text(content)
    lines = content.split("\n")

    for i, line in enumerate(lines, start=1):
        line = line.expandtabs(4)
        if ADD_LINE_NUMBERS:
            add_code_line(doc, f"{i:04d} | {line}")
        else:
            add_code_line(doc, line)


# =========================
# FILE HANDLING
# =========================

def should_skip_dir(path: Path):
    parts = set(path.parts)
    if parts & IGNORED_DIRS:
        return True
    if not INCLUDE_MIGRATIONS and "migrations" in parts:
        return True
    return False


def should_include(path: Path):
    if path.name.startswith("."):
        return False
    if path.suffix.lower() not in ALLOWED_EXTENSIONS:
        return False
    if should_skip_dir(path.parent):
        return False
    try:
        if path.stat().st_size > MAX_FILE_SIZE_KB * 1024:
            return False
    except OSError:
        return False
    return True


def read_file(path: Path):
    encodings = ("utf-8-sig", "utf-8", "utf-16", "cp1251")
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding, errors="strict")
        except UnicodeError:
            continue
        except OSError:
            return None

    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except OSError:
        return None


def sort_key(path: Path, project_root: Path):
    rel = path.relative_to(project_root).as_posix()

    first = rel.split("/", 1)[0]
    try:
        group = APP_ORDER.index(first)
    except ValueError:
        group = len(APP_ORDER)

    return (group, rel)


def collect_files(project_root: Path):
    files = []
    for root, dirs, filenames in os.walk(project_root):
        root_path = Path(root)

        dirs[:] = [
            d for d in dirs
            if not should_skip_dir(root_path / d)
        ]

        for filename in filenames:
            path = root_path / filename
            if should_include(path):
                files.append(path)

    files.sort(key=lambda p: sort_key(p, project_root))
    return files


# =========================
# DESCRIPTIONS — под CityPark
# =========================

def get_description(rel_path: str, content: str):
    rel = rel_path.replace("\\", "/")
    name = Path(rel).name.lower()

    exact = {
        "manage.py": "Точка входа Django-проекта: запуск сервера, выполнение миграций и административных команд.",
        "CityPark/settings.py": "Основные настройки Django-проекта CityPark: приложения, база данных PostgreSQL, шаблоны, статика, медиафайлы и пользовательская модель.",
        "CityPark/urls.py": "Главная маршрутизация проекта: подключение административной панели, страниц сайта, авторизации, корзины, избранного, оформления заказа и профиля.",
        "CityPark/asgi.py": "ASGI-конфигурация проекта для запуска приложения в асинхронной серверной среде.",
        "CityPark/wsgi.py": "WSGI-конфигурация проекта для запуска Django-приложения на веб-сервере.",
        "users/models.py": "Кастомная модель пользователя CustomUser с ФИО, телефоном, email, аватаром, адресом и датами регистрации/обновления.",
        "users/forms.py": "Формы регистрации и авторизации пользователей, включая валидацию логина, ФИО, телефона, email и пароля.",
        "users/views.py": "Представления регистрации и входа пользователя с обработкой форм, сообщениями и перенаправлением на главную страницу.",
        "users/admin.py": "Регистрация и настройка пользовательской модели в административной панели Django.",
        "users/apps.py": "Конфигурация Django-приложения users.",
        "catalog/models.py": "Модели каталога и заказов: категории блюд, блюда, заказы и позиции заказа.",
        "catalog/views.py": "Логика главной страницы, корзины, избранного, профиля, оформления заказа и страницы успешного заказа.",
        "catalog/admin.py": "Настройка отображения категорий, блюд, заказов и позиций заказа в административной панели.",
        "catalog/apps.py": "Конфигурация Django-приложения catalog.",
        "restaurant/models.py": "Модели ресторанного модуля: информация о заведении, столики, бронирования, отзывы, корзина и элементы корзины.",
        "restaurant/views.py": "Представления ресторанного модуля: меню, поиск и сортировка блюд, бронирование столика, контакты, корзина и избранное.",
        "restaurant/urls.py": "Локальные маршруты приложения restaurant для страниц ресторана и пользовательских действий.",
        "restaurant/admin.py": "Настройка административной панели для ресторанных сущностей: бронирований, столиков, отзывов и связанных данных.",
        "restaurant/apps.py": "Конфигурация Django-приложения restaurant.",
        "restaurant/management/commands/fill_db.py": "Пользовательская management-команда для заполнения базы данных начальными данными проекта.",
        "static/js/cart.js": "JavaScript-логика добавления блюд в корзину через AJAX-запрос и получения CSRF-токена из cookie.",
        "README.md": "Краткая документация и название проекта CityPark.",
    }

    if rel in exact:
        return exact[rel]

    template_descriptions = {
        "templates/base.html": "Базовый HTML-шаблон сайта с общей структурой интерфейса, подключением стилей, навигацией и блоком содержимого.",
        "templates/index.html": "Главная страница CityPark с основными блоками сайта и выводом доступных блюд/категорий.",
        "templates/menu.html": "Страница меню с выводом блюд, поиском, фильтрацией, сортировкой и действиями пользователя.",
        "templates/booking.html": "Страница бронирования столика с формой ввода данных посетителя, даты, времени и количества гостей.",
        "templates/about.html": "Информационная страница о заведении CityPark.",
        "templates/contacts.html": "Страница контактов с адресом, телефоном, режимом работы и формой/информационными блоками.",
        "templates/profile.html": "Личный кабинет пользователя с заказами, бронированиями и сводной информацией.",
        "templates/order_success.html": "Страница подтверждения успешного оформления заказа.",
        "templates/basket/basket.html": "Шаблон корзины пользователя с перечнем выбранных блюд, количеством и итоговой стоимостью.",
        "templates/basket/favourites.html": "Шаблон страницы избранных блюд пользователя.",
        "templates/registration/login.html": "Шаблон страницы авторизации пользователя.",
        "templates/registration/register.html": "Шаблон страницы регистрации пользователя.",
    }

    if rel in template_descriptions:
        return template_descriptions[rel]

    if name == "__init__.py":
        return "Служебный файл Python, обозначающий каталог как пакет приложения."
    if name == "tests.py":
        return "Файл для тестов Django-приложения."
    if name.endswith(".html"):
        return "HTML-шаблон интерфейса проекта CityPark."
    if name.endswith(".css"):
        return "Файл стилей пользовательского интерфейса."
    if name.endswith(".js"):
        return "JavaScript-файл клиентской логики сайта."
    if name.endswith(".json"):
        return "JSON-файл данных или конфигурации проекта."
    if name.endswith(".md"):
        return "Markdown-документация проекта."
    if name.endswith(".py"):
        if "class " in content and "models.Model" in content:
            return "Python-модуль с моделями базы данных Django."
        if "def " in content and "render(" in content:
            return "Python-модуль с представлениями Django и обработкой HTTP-запросов."
        return "Python-модуль проекта CityPark."

    return "Файл проекта CityPark."


# =========================
# MAIN
# =========================

def resolve_project_root():
    if len(sys.argv) > 1:
        return Path(sys.argv[1]).resolve()

    cwd = Path.cwd().resolve()
    if (cwd / DEFAULT_PROJECT_DIR).is_dir():
        return (cwd / DEFAULT_PROJECT_DIR).resolve()

    return cwd


def main():
    project_root = resolve_project_root()
    output_path = project_root.parent / OUTPUT_FILE if project_root.name == DEFAULT_PROJECT_DIR else Path.cwd() / OUTPUT_FILE

    if not project_root.exists() or not project_root.is_dir():
        raise FileNotFoundError(f"Папка проекта не найдена: {project_root}")

    files = collect_files(project_root)

    doc = Document()
    set_default_styles(doc)

    # Первая секция: титульный блок листинга в одну колонку
    set_page_layout(doc.sections[0])
    add_page_number(doc.sections[0])

    title = doc.add_heading(f"Листинг кода проекта {PROJECT_NAME}", 0)
    title.alignment = 1

    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run("Автоматически сформированный список исходных файлов проекта")
    r.font.size = Pt(11)

    add_small_paragraph(doc, f"Папка проекта: {project_root.name}", italic=True)
    add_small_paragraph(doc, f"Количество включенных файлов: {len(files)}", italic=True)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # Вторая секция: сам листинг в две колонки
    code_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_page_layout(code_section)
    set_columns(code_section, count=2)
    add_page_number(code_section)

    for index, path in enumerate(files, start=1):
        rel_path = path.relative_to(project_root).as_posix()
        content = read_file(path)

        if content is None:
            continue

        description = get_description(rel_path, content)
        lines_count = len(clean_text(content).splitlines())

        heading = doc.add_heading(f"{index}. {rel_path}", level=2)
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(2)

        add_small_paragraph(doc, f"Назначение файла: {description}")
        add_small_paragraph(doc, f"Количество строк: {lines_count}")

        add_code_block(doc, content)

    doc.save(output_path)
    print(f"Готово: {output_path}")


if __name__ == "__main__":
    main()
